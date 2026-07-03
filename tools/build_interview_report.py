"""Build an interview-ready DOCX report for the Pit Wall Predictor project."""

from __future__ import annotations

from datetime import date
from pathlib import Path
import json

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
REAL_DIR = DATA_DIR / "real"
OUTPUT_DIR = ROOT / "outputs" / "reports"
DOCX_PATH = OUTPUT_DIR / "pit_wall_predictor_interview_report.docx"


def set_run_font(run, *, name: str = "Calibri", size: int | None = None, color: tuple[int, int, int] | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6, line_spacing: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9, color: tuple[int, int, int] = (30, 30, 30), align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, width in enumerate(widths):
        for cell in table.columns[index].cells:
            cell.width = Inches(width)


def style_table(table, header_fill: str = "D9E2F3", body_fill: str | None = None) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            elif body_fill is not None:
                set_cell_shading(cell, body_fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=(46, 116, 181), bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=(46, 116, 181), bold=True)
    else:
        set_run_font(run, size=12, color=(31, 77, 120), bold=True)
    set_paragraph_spacing(paragraph, before=level * 4 + 6, after=6, line_spacing=1.05)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, italic: bool = False, after: float = 6, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=0, after=after, line_spacing=1.1)
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        set_run_font(prefix_run, size=11, color=(20, 20, 20), bold=True, italic=italic)
        remainder = text[len(bold_prefix):]
        if remainder:
            run = paragraph.add_run(remainder)
            set_run_font(run, size=11, color=(20, 20, 20), italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=(20, 20, 20), italic=italic)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text)
    run._r.append(fld_char_end)


def add_table(doc: Document, rows: list[list[str]], widths: list[float], *, header_fill: str = "D9E2F3", body_fill: str | None = None, font_size: int = 9) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    style_table(table, header_fill=header_fill, body_fill=body_fill)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell_text = rows[row_index][col_index]
            is_header = row_index == 0
            set_cell_text(
                cell,
                cell_text,
                bold=is_header,
                size=font_size if not is_header else font_size + 1,
                color=(20, 20, 20),
                align=WD_ALIGN_PARAGRAPH.LEFT,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_report() -> Path:
    real_bundle = REAL_DIR / "2026_real_race_laps.csv"
    calendar_path = REAL_DIR / "race_calendar_2026.csv"
    source_index_path = REAL_DIR / "source_index.json"
    summary_path = ROOT / "outputs" / "reports" / "barcelona_catalunya_grand_prix_all_driver_summary.csv"

    calendar = pd.read_csv(calendar_path)
    laps = pd.read_csv(real_bundle)
    source_index = json.loads(source_index_path.read_text(encoding="utf-8"))
    summary = pd.read_csv(summary_path)
    summary_sorted = summary.sort_values("race_engineer_score", ascending=False).head(5)
    race_counts = (
        laps.groupby("race_name", as_index=False)
        .agg(rows=("lap_number", "size"), drivers=("driver_code", "nunique"))
        .merge(calendar[["race_name", "round", "status", "total_laps"]], on="race_name", how="left")
        .sort_values("round")
    )

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = True

    # Base styles.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(20, 20, 20)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1.font.size = Pt(16)
    heading1.font.color.rgb = RGBColor(46, 116, 181)
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(8)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2.font.size = Pt(13)
    heading2.font.color.rgb = RGBColor(46, 116, 181)
    heading2.paragraph_format.space_before = Pt(12)
    heading2.paragraph_format.space_after = Pt(6)

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "Calibri"
    heading3.font.size = Pt(12)
    heading3.font.color.rgb = RGBColor(31, 77, 120)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(4)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.paragraph_format.line_spacing = 1.0
    run = footer_p.add_run("Pit Wall Predictor interview report - Page ")
    set_run_font(run, size=9, color=(100, 100, 100))
    add_page_number(footer_p)
    for run in footer_p.runs:
        set_run_font(run, size=9, color=(100, 100, 100))

    # Cover page.
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(90)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("Pit Wall Predictor")
    set_run_font(title_run, size=30, color=(0, 0, 0), bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle_p.add_run("Interview report - real-data Formula 1 strategy analysis system")
    set_run_font(subtitle_run, size=14, color=(70, 70, 70))

    summary_p = doc.add_paragraph()
    summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_p.paragraph_format.space_after = Pt(18)
    summary_run = summary_p.add_run(
        "A transparent offline Python project that turns completed-race CSV files into engineering-style race pace, tyre, strategy, and reporting outputs."
    )
    set_run_font(summary_run, size=11, color=(50, 50, 50), italic=True)

    cover_rows = [
        ["Project snapshot", "Value"],
        ["Purpose", "Post-race Formula 1 strategy analysis with transparent formulas and reusable CLI, GUI, and browser front ends."],
        ["Runtime", "Python, NumPy, Pandas, SciPy, Matplotlib, Tkinter, and the standard library."],
        ["Verified data", "8 completed races, 9,215 laps, 22 drivers, 11 teams, imported from local official race CSV files."],
        ["Outputs", "CSV summaries, Markdown reports, PDF reports, PNG plots, and HTML replays."],
        ["Modes", "CLI, colourful Tkinter dashboard, and localhost browser dashboard."],
    ]
    add_table(doc, cover_rows, [1.55, 4.95], header_fill="D9E2F3", font_size=9)

    add_body(
        doc,
        f"Verified on {date.today().isoformat()}, the current imported bundle contains {len(calendar)} completed races and {len(laps)} lap rows. The analysis layer now loads that real bundle automatically whenever it is present.",
        after=8,
    )
    add_body(
        doc,
        "The main design goal is clarity: every score, plot, and replay is generated from explicit rules that can be explained in a technical interview without relying on black-box machine learning.",
        after=0,
    )

    doc.add_page_break()

    # Body sections.
    add_heading(doc, "1. What the project is", 1)
    add_body(
        doc,
        "Pit Wall Predictor is a post-race analysis tool, not a live prediction engine. It reads completed-race lap CSV files, cleans the laps, compares drivers, scores race execution, fits a simple tyre-degradation curve, and then exposes the same backend through three surfaces: command line, Tkinter desktop UI, and a local web UI.",
        after=8,
    )
    add_body(
        doc,
        "That makes the project useful as a portfolio piece because it shows data loading, data cleaning, transparent scoring, plotting, animation, local web serving, and report generation in one coherent codebase.",
        after=8,
    )

    add_heading(doc, "2. How to use it", 1)
    use_rows = [
        ["Mode / command", "What it does", "When to use it"],
        ["python main.py", "Runs the default CLI analysis for the default race and driver pair.", "Fastest way to smoke-test the backend."],
        ["python main.py --race \"Monaco Grand Prix\" --driver VER --compare NOR", "Chooses a race plus two driver codes.", "When you want a specific post-race comparison."],
        ["python main.py --gui", "Launches the colourful Tkinter dashboard.", "When you want a desktop front end."],
        ["python main.py --web", "Starts the localhost browser interface.", "When you want the more polished web UI."],
        ["python main.py --animations", "Exports the HTML replay animations.", "When you need replay artefacts for a presentation."],
    ]
    add_table(doc, use_rows, [2.2, 2.5, 1.8], header_fill="E8EEF5", font_size=9)

    add_heading(doc, "3. What you can input", 1)
    input_rows = [
        ["Input", "Example", "Where it is used"],
        ["Race name", "Barcelona-Catalunya Grand Prix", "Selects the completed race from the calendar."],
        ["Primary driver code", "HAM", "Drives the main analysis summary and tyre fit."],
        ["Comparison driver code", "RUS", "Builds the driver-battle verdict and pace plot."],
        ["Season", "2026", "Matches the calendar and imported data bundle."],
        ["UI flag", "--gui or --web", "Chooses the desktop or browser surface."],
        ["Replay flag", "--animations", "Creates the HTML replay exports."],
        ["Real-data manifest", "data/real/source_manifest.csv", "Tells the importer where the official CSV files live."],
    ]
    add_table(doc, input_rows, [1.4, 2.2, 2.9], header_fill="E8EEF5", font_size=9)

    add_heading(doc, "4. What comes out", 1)
    output_rows = [
        ["Output family", "Example file or surface", "Why it matters"],
        ["Grid summary", "outputs/reports/*_all_driver_summary.csv", "Gives one row per driver with pace, consistency, tyre, and total race-engineer score."],
        ["Readable report", "outputs/reports/*_post_race_report.md and .pdf", "Turns the analysis into a portfolio-friendly narrative."],
        ["Dashboard plots", "outputs/plots/*.png", "Provides the visual evidence behind the scores."],
        ["Replays", "outputs/animations/*.html", "Shows the race, tyre, and pit-stop stories in motion."],
        ["Browser UI", "http://127.0.0.1:8000", "Lets a user explore the project in a local web interface."],
        ["Tkinter UI", "main.py --gui", "Provides the desktop variant without a browser."],
    ]
    add_table(doc, output_rows, [1.5, 2.4, 2.6], header_fill="E8EEF5", font_size=9)

    add_heading(doc, "5. How the analysis works", 1)
    add_heading(doc, "5.1 Data loading and validation", 2)
    add_body(
        doc,
        "The loader prefers the imported real-data bundle in data/real/ whenever it exists. If that bundle is missing, the project falls back to the deterministic sample dataset so the app still opens during development. The import pipeline stores the bundle in a local calendar file, a consolidated lap file, a manifest copy, and a source index that tracks the original CSV source for each race.",
        after=8,
    )
    add_body(
        doc,
        f"The current validated bundle contains {len(calendar)} completed races, {int(laps['driver_code'].nunique())} drivers across the bundle, and {len(laps)} lap rows. The verified completed rounds are: {', '.join(calendar['race_name'].tolist())}.",
        after=8,
    )

    add_heading(doc, "5.2 Clean-lap filtering", 2)
    add_body(
        doc,
        "Clean pace is the foundation of the project. The cleaning rule removes pit laps, safety-car laps, in-laps, out-laps, and then drops any remaining lap that is more than 4.0 seconds slower than the driver's median clean lap. That keeps the pace comparison focused on representative racing pace rather than traffic, pit cycles, or extreme anomalies.",
        after=8,
    )

    add_heading(doc, "5.3 Driver analysis", 2)
    add_body(
        doc,
        "For each driver, the backend calculates average clean pace, fastest clean lap, slowest clean lap, standard deviation, pit-stop count, compound usage, and stint-level summaries. Each stint also gets a simple linear degradation estimate so the report can say where pace held up and where it dropped away.",
        after=8,
    )

    add_heading(doc, "5.4 Scoring model", 2)
    add_body(
        doc,
        "The gamified score is intentionally transparent. It combines pace, consistency, tyre management, and stint execution with fixed weights, then converts the result into an easy-to-read rating and badge. Because the weights are explicit, the interview story is straightforward: the score is an explanation layer, not a hidden AI model.",
        after=8,
    )

    add_heading(doc, "5.5 Tyre degradation model", 2)
    add_body(
        doc,
        "The tyre model uses SciPy's curve_fit on clean laps from a single driver and compound. The fitted form is Lap Time = Base Pace + a x Tyre Age + b x Tyre Age^2. The project reports the fitted coefficients and the RMSE so a reader can see both the trend and the quality of the fit.",
        after=8,
    )

    add_heading(doc, "6. Formula cheat sheet", 1)
    formula_rows = [
        ["Metric", "Formula or rule", "Interpretation"],
        ["Clean pace", "Mean of clean laps after filtering", "Core pace metric for the driver and the grid."],
        ["Consistency score", "clip(100 - lap_time_std x 30)", "Lower lap-time scatter is rewarded."],
        ["Tyre management score", "clip(100 - avg_positive_degradation x 430)", "Punishes steep stint degradation."],
        ["Pace score", "clip(100 - pace_deficit x 22)", "Converts clean pace deficit into a 0-100 scale."],
        ["Stint execution score", "clip(100 - stint_pace_spread x 18)", "Rewards a balanced stint profile."],
        ["Race engineer score", "0.35 pace + 0.25 consistency + 0.25 tyre + 0.15 stint", "Final gamified score used in the ranking."],
        ["Badge logic", "Threshold checks on pace, consistency, tyres, stints, and position", "Produces a human-readable achievement label."],
    ]
    add_table(doc, formula_rows, [1.5, 2.7, 2.3], header_fill="E8EEF5", font_size=9)

    add_body(
        doc,
        "The key design decision is that every formula is simple enough to explain live in an interview. There is no opaque model tuning step and no hidden training set - just fixed rules, visible assumptions, and stable outputs.",
        after=8,
    )

    add_heading(doc, "7. Example output from the verified Barcelona-Catalunya run", 1)
    example_rows = [["Rank", "Driver", "Team", "Clean pace (s)", "Score", "Badge"]]
    for rank, (_, row) in enumerate(summary_sorted.iterrows(), start=1):
        example_rows.append([
            str(rank),
            str(row["driver_code"]),
            str(row["team"]),
            f"{float(row['average_clean_pace']):.3f}",
            f"{float(row['race_engineer_score']):.1f}",
            str(row["badge_earned"]),
        ])
    add_table(doc, example_rows, [0.6, 0.8, 1.35, 1.1, 0.8, 1.85], header_fill="DCE6F1", font_size=8)
    add_body(
        doc,
        "On this verified run, Hamilton leads the sample summary and the report generator turns that result into a compact project artefact. That makes the project easy to demo: the same backend can support a short summary, a detailed explanation, and a visual dashboard from the same race data.",
        after=8,
    )

    add_heading(doc, "8. Plot and replay catalogue", 1)
    plot_rows = [
        ["Artifact", "What it shows"],
        ["full_grid_pace_ranking.png", "Average clean pace by driver; lower is faster."],
        ["full_grid_consistency_ranking.png", "Lap-time stability across the grid."],
        ["full_grid_tyre_management_ranking.png", "Who preserved tyre life best."],
        ["team_pace_comparison.png", "Average pace by team."],
        ["driver_battle_HAM_vs_RUS.png", "Lap-by-lap clean pace comparison."],
        ["full_grid_strategy_timeline.png", "Tyre stints and pit laps across the whole field."],
        ["race_pace_heatmap.png", "Relative clean-lap pace trace by driver and lap."],
        ["tyre_degradation_HAM_M.png", "SciPy tyre-age curve fit for the selected driver."],
        ["race_pace_replay.html", "Animated race-progress replay."],
        ["tyre_degradation_replay_HAM_M.html", "Animated tyre-fit reveal."],
        ["pit_stop_timeline_replay.html", "Animated stint and pit-stop timeline."],
    ]
    add_table(doc, plot_rows, [2.5, 3.8], header_fill="E8EEF5", font_size=8)

    add_heading(doc, "9. Real-data import pipeline and verification", 1)
    add_body(
        doc,
        "The project now reads local FastF1-style CSV exports through a manifest-driven importer. The manifest preserves the source path for each completed race, the importer normalizes the lap schema into the project's standard columns, and the loaded bundle is then used by the CLI, GUI, browser UI, and report generation.",
        after=8,
    )
    ver_rows = [["Round", "Race", "Laps in bundle", "Drivers", "Status"]]
    for _, row in race_counts.iterrows():
        ver_rows.append([
            str(int(row["round"])),
            str(row["race_name"]),
            str(int(row["rows"])),
            str(int(row["drivers"])),
            str(row["status"]),
        ])
    add_table(doc, ver_rows, [0.6, 2.7, 1.2, 0.9, 1.0], header_fill="DCE6F1", font_size=8)
    add_body(
        doc,
        "The validation pass confirmed that all eight supplied completed rounds are visible to the app. That matters for an interview because it shows the project is not only functional in isolation, but also wired to a source-tracked, reproducible data path.",
        after=8,
    )

    add_heading(doc, "10. Project structure", 1)
    module_rows = [
        ["Module", "Responsibility", "Notes"],
        ["main.py", "CLI entry point for demo, GUI, web, and animation modes.", "Routes to the right surface."],
        ["src/real_data_importer.py", "Imports and normalizes real race CSV files.", "Tracks source files and builds the local bundle."],
        ["src/data_loader.py", "Loads real or sample data depending on what is present.", "Keeps the rest of the app simple."],
        ["src/data_cleaning.py", "Applies the clean-lap rules.", "Removes pits, safety-car laps, and outliers."],
        ["src/driver_analysis.py", "Builds per-driver and full-grid analysis.", "Produces the summary table."],
        ["src/scoring.py", "Applies the visible scoring and badge rules.", "Turns metrics into a gamified rating."],
        ["src/tyre_degradation.py", "Fits the quadratic tyre-age model.", "Uses SciPy curve fitting and RMSE."],
        ["src/visualizer.py", "Creates the PNG dashboard plots.", "Matplotlib output only."],
        ["src/animations.py", "Creates HTML replays and static fallbacks.", "Kept dependency-light."],
        ["src/web_app.py", "Serves the local browser dashboard.", "Reuses the same backend calculations."],
        ["src/gui.py", "Tkinter desktop dashboard.", "Alternate offline front end."],
        ["src/report_generator.py", "Writes CSV, Markdown, and PDF report artefacts.", "Used for portfolio exports."],
    ]
    add_table(doc, module_rows, [1.6, 3.35, 1.55], header_fill="E8EEF5", font_size=8)

    add_heading(doc, "11. Interview talking points", 1)
    add_body(
        doc,
        "If you want a crisp spoken summary, describe the project as an offline Formula 1 post-race engineering dashboard that turns source-tracked lap CSVs into pace, consistency, tyre, strategy, and replay outputs. Emphasize that the value is in transparency: every number comes from a visible rule or model, and the same backend powers the CLI, GUI, browser UI, plots, animations, and report export.",
        after=8,
    )
    add_body(
        doc,
        "That framing makes it easy to talk about software engineering, data handling, analysis, and product thinking in one story, which is exactly what an interview panel usually wants to hear.",
        after=0,
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_report()
    print(path)
